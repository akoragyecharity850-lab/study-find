import base64
import zlib
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import io

def encode_kroki(text):
    compressed = zlib.compress(text.encode('utf-8'), 9)
    b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
    return f"https://kroki.io/mermaid/png/{b64}"

arch_diagram = """flowchart TD
    User(["Student/Admin"]) -->|HTTPS| Frontend("React/Vite Frontend")
    Frontend -->|JWT Auth| SupabaseAuth("Supabase Auth")
    Frontend -->|REST API| SupabaseAPI("Supabase PostgREST API")
    SupabaseAPI -->|SQL queries| PostgreSQL[("PostgreSQL Database")]
    SupabaseAuth --> PostgreSQL
    
    subgraph Frontend_Subsystem ["Frontend Subsystem"]
    ReactComponents["React Components"]
    ReactQuery["React Query"]
    Tailwind["Tailwind CSS"]
    ReactComponents --> ReactQuery
    ReactComponents --> Tailwind
    end

    Frontend --- Frontend_Subsystem
    
    subgraph Backend_Subsystem ["Backend Subsystem"]
    SupabaseAuth
    SupabaseAPI
    PostgreSQL
    end
"""

er_diagram = """erDiagram
    PROFILES {
        uuid id PK
        string full_name
        string program
        int year
    }
    STUDY_GROUPS {
        uuid id PK
        string name
        string course
        uuid leader_id FK
    }
    MEMBERS {
        uuid id PK
        uuid group_id FK
        uuid user_id FK
    }
    POSTS {
        uuid id PK
        uuid group_id FK
        uuid user_id FK
        string content
    }
    SESSIONS {
        uuid id PK
        uuid group_id FK
        string title
        date session_date
    }

    PROFILES ||--o{ STUDY_GROUPS : "leads"
    PROFILES ||--o{ MEMBERS : "in"
    STUDY_GROUPS ||--o{ MEMBERS : "has"
    PROFILES ||--o{ POSTS : "writes"
    STUDY_GROUPS ||--o{ POSTS : "contains"
    PROFILES ||--o{ SESSIONS : "makes"
    STUDY_GROUPS ||--o{ SESSIONS : "sets"
"""

def download_image(url):
    response = requests.get(url)
    if response.status_code != 200:
        print(f"Failed to download image: {response.status_code}")
        print(response.text)
        sys.exit(1)
    return io.BytesIO(response.content)

print("Downloading Architecture Diagram...")
arch_url = encode_kroki(arch_diagram)
arch_img = download_image(arch_url)

print("Downloading ER Diagram...")
er_url = encode_kroki(er_diagram)
er_img = download_image(er_url)

print("Building Document...")
document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Trebuchet MS'
font.size = Pt(12)

def add_heading(text, level, align=WD_ALIGN_PARAGRAPH.LEFT):
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Trebuchet MS'
    heading.alignment = align

def add_paragraph(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Trebuchet MS'
    run.font.size = Pt(12)
    run.bold = bold
    return p

# Cover Page
add_heading('UGANDA CHRISTIAN UNIVERSITY', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading('FACULTY OF ENGINEERING, DESIGN AND TECHNOLOGY', level=2, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('')
add_paragraph('COURSE CODE: CSC1202', align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('COURSE NAME: WEB AND MOBILE APPLICATION DEVELOPMENT', align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('')
add_heading('PROJECT REPORT', level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading('"STUDENT STUDY GROUP FINDER"', level=2, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('')
add_paragraph('PROGRAM: BACHELOR OF SCIENCE IN INFORMATION TECHNOLOGY (BSIT) Y1S2', align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph('SUBMISSION DATE: APRIL 2026', align=WD_ALIGN_PARAGRAPH.CENTER)
document.add_page_break()

add_heading('1. Introduction', level=2)
add_paragraph('Collaborative learning is a critical aspect of tertiary education, providing students the platform to enhance their academic experiences. Currently, finding and organizing study groups at Uganda Christian University largely depends on informal personal networks and impromptu social media interactions, often limiting broader participation.')
add_paragraph('The Student Study Group Finder aims to directly address this limitation. It provides a structured, digital platform where students can create, discover, and actively manage academic study groups. This ensures a centralized, accessible, and inclusive environment, promoting academic excellence and collaborative learning across all university faculties.')

add_heading('2. Problem Statement & Objectives', level=2)
add_heading('Problem Statement', level=3)
add_paragraph('The challenge students face is the fragmented and isolated nature of forming academic alliances. Without a centralized hub, students missing initial social connections struggle to find study groups tailored to their specific courses, minimizing equal learning opportunities.')
add_heading('Objectives', level=3)
p = add_paragraph('1. To implement secure user authentication and student profile management.\n2. To provide functionalities for the creation, management, and discovery of study groups based on courses and faculties.\n3. To enable group leaders to schedule and organize study sessions seamlessly.\n4. To implement group communication tools (posts/announcements) for collaborative preparation.')

add_heading('3. Technology Stack & Choices', level=2)
add_paragraph('In line with fulfilling the requirement of a scalable and robust minimal viable product (MVP), the following stack was utilized:\n- Frontend Framework: React.js configured with Vite. Selected for fast rendering, component reusability, and efficient state management.\n- Styling: Tailwind CSS & Radix UI (shadcn/ui). Chosen for utility-first styling and extremely accessible User Interface components, ensuring a robust responsive design on desktop and mobile.\n- Backend Framework / Database: Supabase (PostgreSQL with PostgREST). Chosen for secure JWT authentication out-of-the-box, real-time table capabilities, and reliable relational configurations mimicking standard backend logic.\n- Hosting / Deployment: GitHub Pages (frontend delivery).')

add_heading('4. System Architecture', level=2)
add_paragraph('The application leverages a modern Client-Server Architecture utilizing a Backend-as-a-Service (BaaS) infrastructure. This perfectly maximizes UI rendering efficiency by offloading work to the client, while abstracting and standardizing backend API integrations smoothly.')

p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(arch_img, width=Inches(5.5))
add_paragraph('Figure 1: Client-Server Architecture utilizing BaaS Component structure.', align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading('5. Entity Relationship (ER) Diagram', level=2)
add_paragraph('The database relies on a highly normalized relational setup leveraging PostgreSQL types to preserve referential integrity across modules.')

p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(er_img, width=Inches(4.5))
add_paragraph('Figure 2: PostgreSQL Entity-Relationship Schema mapping.', align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading('6. System Features & User Manual', level=2)
add_heading('6.1 Registration and Authentication', level=3)
add_paragraph('- Step 1: Navigate to the web application URL.\n- Step 2: Click on Sign Up / Login. Input your correct Full Name, Program of Study, Year of Study, institutional Email Address, and secure Password.\n- Step 3: Once successfully verified, visually seamlessly navigate via the populated Dashboards.')
add_heading('6.2 Discovering and Managing Study Groups', level=3)
add_paragraph('- Finding a Group: Under the groups interface, browse the active inventory of academic teams. Utilizing the Search and Filter fields, drill down by specific Course Name or respective Faculty.\n- Joining: Verify a group’s intent from its description then click "Join Group" to gain timeline access.\n- Creating a Group: In "Create Module", submit crucial configurations (Group Name, Course Code, Description, Location Format). Note: The creator automatically holds default Group Leader privileges.')
add_heading('6.3 Group Hub: Communication and Scheduling', level=3)
add_paragraph('- Scheduling Sessions (Leaders Only): Select "Create Session" in the Hub interface. Establish the Date, precise Time, Physical Location or Virtual Link, and a concise outline of modules to cover. These immediately reflect on member dashboards.\n- Interactions: The built-in "Discussion Wall" mechanism tracks inquiries, shared resource linkages, and casual study notifications in a single persistent timeline.')

add_heading('7. API Documentation', level=2)
add_paragraph('The Application Programming Interfaces strictly adhere to well-formatted REST standard methodologies securely orchestrated by PostgREST acting over PostgreSQL, returning normalized JSON payload objects and arrays.')
add_paragraph("Endpoints & Functions:\n1. POST /auth/v1/signup: Student Registration Identity creation.\n2. POST /auth/v1/token?grant_type=password: Authentication token exchange.\n3. GET /rest/v1/study_groups: Fetch groups; manages filtering queries.\n4. POST /rest/v1/study_groups: Create new study ecosystem. (Valid JWT header expected).\n5. POST /rest/v1/study_sessions: Write new group study session record entry.\n6. GET /rest/v1/group_posts?group_id=eq.[id]: Fetches specific ongoing discussion pipelines.")

add_heading('8. GitHub Collaborative Deliverables', level=2)
add_paragraph('The core repositories encompassing source code and database rules/schemas conform to best practices highlighting detailed commitment traces and collaborative milestones.')
add_paragraph('- Application Main Repository Filepath (Frontend Logic): https://github.com/evarduwamani-ctrl/Student-Study-Group-Finder\n- Backend Environment Pipeline: Systematized and tightly coupled actively running via Supabase integration frameworks. Authentication logic and relational setup exist organically on the database node structure driven identically by the primary repository connections.')

add_heading('9. References', level=2)
add_paragraph('[1] D. Flanagan, JavaScript: The Definitive Guide. 7th ed. Sebastopol, CA: O\'Reilly Media, 2020.\n[2] "Supabase Official Documentation," Supabase HQ. [Online]. Available: https://supabase.com/docs.\n[3] "React - A JavaScript library for building user interfaces," Meta Platforms. [Online]. Available: https://reactjs.org/.\n[4] "Tailwind CSS - Rapidly build modern websites without ever leaving your HTML," Tailwind Labs. [Online]. Available: https://tailwindcss.com/.\n[5] "IEEE Reference Guide," IEEE, Piscataway, NJ, USA, 2018.')

document.save('Project_Report.docx')
print("Document generated successfully saved as Project_Report.docx")
